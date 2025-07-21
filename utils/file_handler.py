import os
from PyPDF2 import PdfReader
from docx import Document

def load_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def load_pdf(file_path):
    text = ""
    with open(file_path, "rb") as f:
        reader = PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text

def load_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def load_files(folder_path):
    from langchain_core.documents import Document as LCDocument
    docs = []
    for file in os.listdir(folder_path):
        full_path = os.path.join(folder_path, file)
        if file.endswith(".txt"):
            content = load_txt(full_path)
        elif file.endswith(".pdf"):
            content = load_pdf(full_path)
        elif file.endswith(".docx"):
            content = load_docx(full_path)
        else:
            continue
        docs.append(LCDocument(page_content=content, metadata={"source": file}))
    return docs